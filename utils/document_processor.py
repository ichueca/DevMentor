"""
Procesador de Documentos para RAG

Soporta múltiples formatos
- PDF (.pdf)
- Word (.docx)
- Texto plano (.txt, .md)

Extrae texto limpio y validado
"""

import os
import logging
from pathlib import Path
from typing import Tuple, Optional

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """ Procesa documentos en múltiples formatos """

    SUPPORTED_FORMATS = {'.pdf','.docx', '.txt', '.md'}
    # 100 MB
    MAX_FILE_SIZE = 100 * 1024 * 1024 

    @staticmethod
    def process_document(file_path: str) -> Tuple[str,str]:
        """
        Procesa un documento y extrae su texto

        Args:
            file_path: La ruta al documento

        Returns:
            Tupla[texto extraido, nombre documento]

        Raises:
            ValueError: Si el formato no está soportado o hay errores
            FileNotFoundError: Si el documento no existe
        """
        file_Path = Path(file_path)

        if not file_Path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        file_ext = file_Path.suffix.lower()
        if file_ext not in DocumentProcessor.SUPPORTED_FORMATS:
            raise ValueError(f"Formato no soportado: {file_ext}  | Soportados: {DocumentProcessor.SUPPORTED_FORMATS}")

        file_size = file_Path.stat().st_size
        if file_size > DocumentProcessor.MAX_FILE_SIZE:
            raise ValueError(f"Archivo excede de {DocumentProcessor.MAX_FILE_SIZE/1024/1024:.1f} MB ({file_size/1024/1024:.1f} MB)")
        
        if file_ext == '.pdf':
            text = DocumentProcessor._extract_pdf(file_Path)
        elif file_ext == '.docx':
            text = DocumentProcessor._extract_docx(file_Path)
        else:
            text = DocumentProcessor._extract_text(file_Path)

        if not text or not text.strip():
            raise ValueError(f"No se pudo extraer texto del archuvo: {file_Path.name}")

        logger.info(f"Texto extraído: {len(text)} caracteres")
        return text, file_Path.name
    
    @staticmethod
    def _extract_pdf(file_path:Path) -> str:
        """ Extrae el texto de un PDF """
        try:
            reader = PdfReader(file_path)
            text = ""
            for page_num, page in enumerate(reader.pages, 1):
                text += f"\n--- Página {page_num} ---\n"
                text += page.extract_text()
            return text
        except Exception as e:
            raise ValueError(f"Error al procesar {file_path.name} : {str(e)}")
    
    @staticmethod
    def _extract_docx(file_path:Path) -> str:
        """ Extrae el texto de un documento de Word (docx """
        try:
            doc = Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text
            for table in doc.tables:
                text += "\n-- Tabla --\n"
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells
                    )
                    text += row_text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error al procesar {file_path.name} : {str(e)}")
    
    @staticmethod
    def _extract_text(file_path:Path) -> str:
        """ Extrae el texto de un documento de texto plano """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-8') as f:
                    return f.read()
            except Exception as e:
                raise ValueError(f"Error al procesar {file_path.name} : {str(e)}")
        except Exception as e:
            raise ValueError(f"Error al procesar {file_path.name} : {str(e)}")
        

if __name__ == '__main__':
    import sys
    from pathlib import Path

    print("\n" + "="*70)
    print("🧪 PRUEBA DE DOCUMENT PROCESSOR")
    print("="*70)

    doc_folder = Path("documents")

    documents = list(doc_folder.glob("*"))

    if not documents:
        print(f"❌ No hay documentos en la carpeta")
        sys.exit(1)

    print(f"📁 Documentos encontrados: {len(documents)}\n")

    results[]
    for doc_path in sorted(documents):
        if doc_path.is_file():
            print(f"📄 Procesando {doc_path.name}")
            print("-"*70)

            try:
                text, filename = DocumentProcessor.process_document(str(doc_path))
                